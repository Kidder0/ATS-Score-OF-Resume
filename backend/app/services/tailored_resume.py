import re
from io import BytesIO
from typing import Any

from docx import Document

from app.services.grounding import (
    GUARDRAIL_WARNING,
    build_professional_summary,
    classify_keyword_support,
    rewrite_resume_bullets,
)
from app.services.keyword_extractor import extract_job_keywords
from app.services.scoring import score_resume_match
from app.services.text_utils import sentences, unique_preserve_order


def build_tailored_resume(resume_text: str, job_description: str) -> dict[str, Any]:
    keywords = extract_job_keywords(job_description)
    score, _, matched, missing, weak, weak_areas = score_resume_match(resume_text, keywords)
    bullets, genuine_gaps = rewrite_resume_bullets(resume_text, keywords, max_bullets=8)
    supported_keywords, unsupported_keywords = classify_keyword_support(
        resume_text,
        unique_preserve_order(keywords.skills + keywords.tools),
    )
    name = _candidate_name(resume_text)
    contact_lines = _contact_lines(resume_text)
    project_lines = _project_or_experience_lines(resume_text)
    summary = build_professional_summary(resume_text, keywords)
    gap_plan = _gap_action_plan(unique_preserve_order(genuine_gaps + unsupported_keywords), weak)
    section_suggestions = _build_section_suggestions(
        resume_text=resume_text,
        summary=summary,
        supported_keywords=supported_keywords,
        rewritten_bullets=bullets,
        education_lines=_education_lines(resume_text),
        gap_plan=gap_plan,
    )

    return {
        "name": name,
        "contact_lines": contact_lines,
        "target_title": "Entry-Level AI Engineer / GenAI Developer",
        "summary": summary,
        "supported_keywords": supported_keywords[:18],
        "rewritten_bullets": bullets,
        "project_or_experience_lines": project_lines[:8],
        "education_lines": _education_lines(resume_text),
        "match_score": score,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "weak_keywords": weak,
        "weak_areas": weak_areas,
        "genuine_gaps": unique_preserve_order(genuine_gaps + unsupported_keywords),
        "guardrail_warning": GUARDRAIL_WARNING,
        "gap_action_plan": gap_plan,
        "section_suggestions": section_suggestions,
        "preserved_format_note": (
            "Suggestions are grouped by the uploaded resume's existing section headings. "
            "Unsupported gaps are not inserted into resume sections."
        ),
    }


def render_tailored_resume_markdown(resume: dict[str, Any]) -> str:
    lines = [
        f"# {resume['name']}",
        "",
        *resume.get("contact_lines", []),
        "",
    ]
    for section in resume.get("section_suggestions", []):
        if section.get("is_gap_plan"):
            continue
        lines.extend(["", f"## {section['section_title']}"])
        for item in section.get("suggested_lines", []):
            lines.append(f"- {item}" if section.get("as_bullets") else str(item))
    lines.extend(
        [
            "",
            "---",
            "",
            "## Truthfulness Review",
            resume["guardrail_warning"],
            "",
            "## Genuine Gaps Not Added To Resume",
            _markdown_list(resume.get("gap_action_plan", [])),
        ]
    )
    return "\n".join(line for line in lines if line is not None)


def render_tailored_resume_docx(resume: dict[str, Any]) -> BytesIO:
    document = Document()
    document.add_heading(resume["name"], level=1)
    for line in resume.get("contact_lines", []):
        document.add_paragraph(line)
    for section in resume.get("section_suggestions", []):
        if section.get("is_gap_plan"):
            continue
        document.add_heading(section["section_title"], level=2)
        style = "List Bullet" if section.get("as_bullets") else None
        for line in section.get("suggested_lines", []):
            document.add_paragraph(str(line), style=style)
    document.add_page_break()
    document.add_heading("Truthfulness Review", level=2)
    document.add_paragraph(resume["guardrail_warning"])
    document.add_heading("Genuine Gaps Not Added To Resume", level=2)
    for item in resume.get("gap_action_plan", []):
        document.add_paragraph(item, style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _candidate_name(resume_text: str) -> str:
    for line in resume_text.splitlines():
        clean = line.strip(" #\t")
        if not clean:
            continue
        if clean.startswith("-") or "," in clean:
            continue
        if re.search(r"resume|engineer|developer|summary|skills|projects|education", clean, re.I):
            continue
        if 1 <= len(clean.split()) <= 5:
            return clean
        break
    return "Candidate"


def _contact_lines(resume_text: str) -> list[str]:
    output: list[str] = []
    for line in resume_text.splitlines()[:12]:
        if re.search(r"@|linkedin|github|https?://|\(?\d{3}\)?[-.\s]\d{3}[-.\s]\d{4}", line, re.I):
            output.append(line.strip())
    return unique_preserve_order(output)


def _project_or_experience_lines(resume_text: str) -> list[str]:
    return [
        _clean_line(sentence)
        for sentence in sentences(resume_text)
        if re.search(r"\b(built|developed|implemented|created|designed|deployed|analyzed|automated)\b", sentence, re.I)
    ]


def _education_lines(resume_text: str) -> list[str]:
    output: list[str] = []
    for line in resume_text.splitlines():
        clean = _clean_line(line)
        if len(clean) > 180:
            continue
        if re.search(r"\b(university|college|school|bachelor|master|ms in|be in|bs in|degree|gpa)\b", clean, re.I):
            output.append(clean)
    return unique_preserve_order(output)[:4]


def _gap_action_plan(gaps: list[str], weak_keywords: list[str]) -> list[str]:
    if not gaps and not weak_keywords:
        return ["No unsupported gaps were detected for the extracted job signals."]
    actions: list[str] = []
    for keyword in weak_keywords[:6]:
        actions.append(f"Strengthen '{keyword}' only by tying it to an existing project, tool, course, or measurable outcome already supported by the resume.")
    for keyword in gaps[:10]:
        actions.append(f"Do not add '{keyword}' to the resume until you have real experience, coursework, certification, or a project that supports it.")
    return unique_preserve_order(actions)


def _markdown_list(items: list[str]) -> str:
    if not items:
        return "- None supported by the uploaded resume"
    return "\n".join(f"- {item}" for item in items)


def _clean_line(line: str) -> str:
    return re.sub(r"\s+", " ", line.strip(" -*\t")).rstrip(".")


def _build_section_suggestions(
    resume_text: str,
    summary: str,
    supported_keywords: list[str],
    rewritten_bullets: list[str],
    education_lines: list[str],
    gap_plan: list[str],
) -> list[dict[str, Any]]:
    sections = _parse_resume_sections(resume_text)
    if not sections:
        return [
            {
                "section_title": "Professional Summary",
                "suggested_lines": [summary],
                "as_bullets": False,
                "note": "Added because no clear section headings were detected.",
            },
            {
                "section_title": "Skills",
                "suggested_lines": supported_keywords[:18],
                "as_bullets": True,
                "note": "Only resume-supported JD keywords are included.",
            },
            {
                "section_title": "Projects / Experience",
                "suggested_lines": rewritten_bullets,
                "as_bullets": True,
                "note": "Rewritten from uploaded resume evidence.",
            },
        ]

    suggestions: list[dict[str, Any]] = []
    used_summary = False
    used_skills = False
    used_bullets = False
    used_education = False
    for section in sections:
        title = section["title"]
        kind = _section_kind(title)
        original_lines = section["lines"]
        suggested_lines = original_lines[:]
        note = "Original section preserved."
        as_bullets = _section_uses_bullets(original_lines)

        if kind == "summary":
            suggested_lines = [summary]
            used_summary = True
            as_bullets = False
            note = "Summary rewritten using resume-supported JD keywords only."
        elif kind == "skills":
            suggested_lines = _merge_skill_lines(original_lines, supported_keywords)
            used_skills = True
            as_bullets = True
            note = "Supported JD keywords are suggested inside the existing skills section."
        elif kind == "experience":
            suggested_lines = rewritten_bullets or original_lines
            used_bullets = True
            as_bullets = True
            note = "Bullets rewritten from uploaded resume evidence."
        elif kind == "education":
            suggested_lines = education_lines or original_lines
            used_education = True
            as_bullets = True
            note = "Education section preserved; missing certifications are not invented."

        suggestions.append(
            {
                "section_title": title,
                "original_lines": original_lines,
                "suggested_lines": unique_preserve_order(suggested_lines),
                "as_bullets": as_bullets,
                "note": note,
            }
        )

    if not used_summary:
        suggestions.insert(
            0,
            {
                "section_title": "Professional Summary",
                "original_lines": [],
                "suggested_lines": [summary],
                "as_bullets": False,
                "note": "Suggested because no summary/profile section was detected.",
            },
        )
    if supported_keywords and not used_skills:
        suggestions.append(
            {
                "section_title": "Skills",
                "original_lines": [],
                "suggested_lines": supported_keywords[:18],
                "as_bullets": True,
                "note": "Suggested because supported JD keywords were found but no skills section was detected.",
            }
        )
    if rewritten_bullets and not used_bullets:
        suggestions.append(
            {
                "section_title": "Projects / Experience",
                "original_lines": [],
                "suggested_lines": rewritten_bullets,
                "as_bullets": True,
                "note": "Suggested from uploaded resume evidence because no project/experience section was detected.",
            }
        )
    if education_lines and not used_education:
        suggestions.append(
            {
                "section_title": "Education",
                "original_lines": [],
                "suggested_lines": education_lines,
                "as_bullets": True,
                "note": "Education detected from the uploaded resume.",
            }
        )

    suggestions.append(
        {
            "section_title": "Gap Action Plan - Not Added To Resume",
            "original_lines": [],
            "suggested_lines": gap_plan,
            "as_bullets": True,
            "note": "These gaps are not inserted into resume sections because the uploaded resume does not support them.",
            "is_gap_plan": True,
        }
    )
    return suggestions


def _parse_resume_sections(resume_text: str) -> list[dict[str, Any]]:
    lines = [_clean_line(line) for line in resume_text.splitlines()]
    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for line in lines:
        if not line:
            continue
        if _is_heading(line):
            current = {"title": line.strip("# ").title(), "lines": []}
            sections.append(current)
            continue
        if current is not None:
            current["lines"].append(line)
    return [section for section in sections if section["lines"] or _section_kind(section["title"]) != "other"]


def _is_heading(line: str) -> bool:
    clean = line.strip("# ").strip()
    if not clean or len(clean) > 60:
        return False
    known = r"(summary|profile|objective|skills|technical skills|experience|work experience|projects|project experience|education|certifications|certification)"
    if re.fullmatch(known, clean, re.I):
        return True
    return clean.isupper() and len(clean.split()) <= 5


def _section_kind(title: str) -> str:
    clean = title.lower()
    if any(word in clean for word in ["summary", "profile", "objective"]):
        return "summary"
    if "skill" in clean or "technical" in clean:
        return "skills"
    if any(word in clean for word in ["experience", "project", "work"]):
        return "experience"
    if "education" in clean or "degree" in clean:
        return "education"
    return "other"


def _section_uses_bullets(lines: list[str]) -> bool:
    return bool(lines) and (len(lines) > 1 or any(line.startswith("-") for line in lines))


def _merge_skill_lines(original_lines: list[str], supported_keywords: list[str]) -> list[str]:
    existing = ", ".join(original_lines)
    additions = [keyword for keyword in supported_keywords if not re.search(rf"(?<!\w){re.escape(keyword)}(?!\w)", existing, re.I)]
    if not additions:
        return original_lines
    return original_lines + [", ".join(additions[:12])]
