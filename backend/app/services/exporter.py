from io import BytesIO
from typing import Any

from docx import Document


def render_markdown_report(report: dict[str, Any]) -> str:
    lines = [
        "# AI Resume & Job Match Report",
        "",
        f"**Match score:** {report.get('match_score', 'N/A')}%",
        f"**Fit level:** {report.get('fit_level', 'N/A')}",
        "",
        "## Guardrail",
        str(report.get("guardrail_warning", "Do not invent experience.")),
        "",
        "## Score Breakdown",
    ]
    for item in report.get("score_breakdown", []):
        lines.extend(
            [
                f"- **{item.get('category')} ({item.get('weight')}%)**: {item.get('earned_points')} points",
                f"  - {item.get('rationale')}",
            ]
        )
    lines.extend(["", "## Matched Keywords", _list(report.get("matched_keywords", []))])
    lines.extend(["", "## Missing Keywords", _list(report.get("missing_keywords", []))])
    lines.extend(["", "## Weak Areas", _list(report.get("weak_areas", []))])
    lines.extend(["", "## Evidence Map"])
    for item in report.get("evidence", []):
        status = item.get("status", "unknown")
        evidence = item.get("evidence") or item.get("recommendation", "")
        lines.append(f"- **{item.get('keyword')}** ({status}): {evidence}")
    lines.extend(["", "## Tailored Summary", str(report.get("rewritten_summary", ""))])
    lines.extend(["", "## Rewritten Bullets", _list(report.get("rewritten_bullets", []))])
    lines.extend(["", "## Cover Letter", str(report.get("cover_letter", ""))])
    lines.extend(["", "## Recommendations", _list(report.get("recommendations", []))])
    return "\n".join(lines)


def render_docx_report(report: dict[str, Any]) -> BytesIO:
    document = Document()
    document.add_heading("AI Resume & Job Match Report", level=1)
    document.add_paragraph(f"Match score: {report.get('match_score', 'N/A')}%")
    document.add_paragraph(f"Fit level: {report.get('fit_level', 'N/A')}")
    document.add_heading("Guardrail", level=2)
    document.add_paragraph(str(report.get("guardrail_warning", "Do not invent experience.")))
    document.add_heading("Score Breakdown", level=2)
    for item in report.get("score_breakdown", []):
        document.add_paragraph(
            f"{item.get('category')} ({item.get('weight')}%): {item.get('earned_points')} points",
            style="List Bullet",
        )
        document.add_paragraph(str(item.get("rationale", "")))
    for heading, key in [
        ("Matched Keywords", "matched_keywords"),
        ("Missing Keywords", "missing_keywords"),
        ("Weak Areas", "weak_areas"),
        ("Rewritten Bullets", "rewritten_bullets"),
        ("Recommendations", "recommendations"),
    ]:
        document.add_heading(heading, level=2)
        for value in report.get(key, []):
            document.add_paragraph(str(value), style="List Bullet")
    document.add_heading("Evidence Map", level=2)
    for item in report.get("evidence", []):
        document.add_paragraph(
            f"{item.get('keyword')} ({item.get('status')}): {item.get('evidence') or item.get('recommendation', '')}",
            style="List Bullet",
        )
    document.add_heading("Tailored Summary", level=2)
    document.add_paragraph(str(report.get("rewritten_summary", "")))
    document.add_heading("Cover Letter", level=2)
    document.add_paragraph(str(report.get("cover_letter", "")))
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _list(items: list[Any]) -> str:
    if not items:
        return "- None detected"
    return "\n".join(f"- {item}" for item in items)
