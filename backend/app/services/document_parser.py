from pathlib import Path

from docx import Document
from pypdf import PdfReader


ALLOWED_EXTENSIONS = {".pdf", ".docx"}


class DocumentParseError(ValueError):
    pass


def validate_upload(filename: str, content_type: str | None, file_size: int, max_size: int) -> None:
    extension = Path(filename).suffix.lower()
    allowed_content = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentParseError("Only PDF and DOCX resume uploads are supported.")
    if content_type and content_type not in allowed_content:
        raise DocumentParseError("Uploaded file type does not match PDF or DOCX.")
    if file_size > max_size:
        raise DocumentParseError("Resume file is too large.")


def extract_resume_text(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".pdf":
        text = _extract_pdf_text(path)
    elif extension == ".docx":
        text = _extract_docx_text(path)
    else:
        raise DocumentParseError("Unsupported resume format.")

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(normalized) < 40:
        raise DocumentParseError("Could not extract enough text from the resume.")
    return normalized


def _extract_pdf_text(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse PDF resume: {exc}") from exc


def _extract_docx_text(path: Path) -> str:
    try:
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise DocumentParseError(f"Failed to parse DOCX resume: {exc}") from exc

