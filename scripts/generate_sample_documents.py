from pathlib import Path

from docx import Document
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
RESUME_MD = ROOT / "samples" / "resumes" / "entry_level_ai_engineer_resume.md"
OUTPUT_DIR = ROOT / "samples" / "resumes"


def main() -> None:
    text = RESUME_MD.read_text(encoding="utf-8")
    create_docx(text, OUTPUT_DIR / "entry_level_ai_engineer_resume.docx")
    create_pdf(text, OUTPUT_DIR / "entry_level_ai_engineer_resume.pdf")


def create_docx(text: str, path: Path) -> None:
    document = Document()
    for line in text.splitlines():
        if line.startswith("# "):
            document.add_heading(line.removeprefix("# "), level=1)
        elif line.startswith("## "):
            document.add_heading(line.removeprefix("## "), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line.removeprefix("- "), style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    document.save(path)


def create_pdf(text: str, path: Path) -> None:
    pdf = canvas.Canvas(str(path))
    width, height = 612, 792
    y = height - 72
    for line in text.splitlines():
        if not line.strip():
            y -= 12
            continue
        pdf.drawString(72, y, line[:95])
        y -= 16
        if y < 72:
            pdf.showPage()
            y = height - 72
    pdf.save()


if __name__ == "__main__":
    main()
